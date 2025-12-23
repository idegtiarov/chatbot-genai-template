"""Utility module collecting file parsing functions for RAG documents"""

import io
from logging import getLogger

logger = getLogger(__name__)


def parse_text_file(content: bytes) -> str:
    """Parse plain text file with automatic encoding detection"""
    for encoding in ["utf-8", "utf-16", "latin-1", "cp1252"]:
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue

    raise ValueError("Could not decode text file with any supported encoding")


def parse_pdf_file(content: bytes) -> str:
    """Parse PDF file and extract text"""
    try:
        import pypdf
    except ImportError as e:
        raise ValueError("PDF parsing requires pypdf package") from e

    try:
        pdf_reader = pypdf.PdfReader(io.BytesIO(content))
        text_parts = [page.extract_text() for page in pdf_reader.pages if page.extract_text()]
        return "\n\n".join(text_parts)

    except Exception as e:
        logger.error("Failed to parse PDF: %s", e)
        raise ValueError(f"Failed to parse PDF file: {e}") from e


def parse_docx_file(content: bytes) -> str:
    """Parse Word document and extract text"""
    try:
        import docx
    except ImportError as e:
        raise ValueError("DOCX parsing requires python-docx package") from e

    try:
        doc = docx.Document(io.BytesIO(content))
        text_parts = []

        # Extract paragraphs
        text_parts.extend(p.text for p in doc.paragraphs if p.text.strip())

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        return "\n\n".join(text_parts)

    except Exception as e:
        logger.error("Failed to parse DOCX: %s", e)
        raise ValueError(f"Failed to parse DOCX file: {e}") from e


PARSER_MAP = {
    ".txt": parse_text_file,
    ".pdf": parse_pdf_file,
    ".docx": parse_docx_file,
}
