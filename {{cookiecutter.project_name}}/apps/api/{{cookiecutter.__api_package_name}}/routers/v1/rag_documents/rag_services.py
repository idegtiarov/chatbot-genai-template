"""File parsing and validation services for RAG document upload"""

import io
from dataclasses import dataclass
from logging import getLogger

from fastapi import UploadFile

from ....app.settings import settings

logger = getLogger(__name__)


class FileValidationError(Exception):
    """Raised when file validation fails"""

    def __init__(self, message: str, status_code: int = 400):
        super().__init__(message)
        self.message = message
        self.status_code = status_code


@dataclass
class ParsedFile:
    """Result of file validation and parsing"""

    content: bytes
    text_content: str
    filename: str
    content_type: str
    file_ext: str


def _get_file_extension(filename: str) -> str:
    """Extract lowercase file extension from filename (e.g., '.pdf')"""
    return "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""


async def validate_and_parse_file(file: UploadFile) -> ParsedFile:
    """
    Validate and parse an uploaded file.

    Performs:
    - File extension validation
    - File size validation
    - Text content extraction

    Args:
        file: The uploaded file

    Returns:
        ParsedFile with validated content and extracted text

    Raises:
        FileValidationError: If validation fails
    """
    filename = file.filename or "unknown"
    file_ext = _get_file_extension(filename)

    # Validate file extension
    allowed_extensions = settings.RAG.allowed_extensions.split(",")
    if file_ext not in allowed_extensions:
        raise FileValidationError(
            f"File type not allowed. Allowed types: {', '.join(allowed_extensions)}",
            status_code=400,
        )

    # Read and validate file size
    content = await file.read()
    max_size = settings.RAG.max_file_size_mb * 1024 * 1024
    if len(content) > max_size:
        raise FileValidationError(
            f"File too large. Maximum size: {settings.RAG.max_file_size_mb}MB",
            status_code=413,
        )

    # Extract text content
    content_type = file.content_type or ""
    text_content = await _parse_file_content(content, file_ext, content_type)

    if not text_content.strip():
        raise FileValidationError("Could not extract text content from file")

    return ParsedFile(
        content=content,
        text_content=text_content,
        filename=filename,
        content_type=content_type,
        file_ext=file_ext,
    )


async def _parse_file_content(content: bytes, file_ext: str, content_type: str) -> str:
    """
    Extract text content from file bytes.

    Args:
        content: Raw file bytes
        file_ext: File extension (e.g., '.pdf')
        content_type: MIME type of the file

    Returns:
        Extracted text content

    Raises:
        ValueError: If file type is not supported or parsing fails
    """
    if file_ext == ".txt" or content_type == "text/plain":
        return _parse_text_file(content)

    if file_ext == ".pdf" or content_type == "application/pdf":
        return await _parse_pdf_file(content)

    if file_ext in (".docx", ".doc") or content_type in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ):
        return await _parse_docx_file(content, file_ext)

    raise ValueError(f"Unsupported file type: {file_ext or content_type}")


def _parse_text_file(content: bytes) -> str:
    """Parse plain text file with automatic encoding detection"""
    for encoding in ["utf-8", "utf-16", "latin-1", "cp1252"]:
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue

    raise ValueError("Could not decode text file with any supported encoding")


async def _parse_pdf_file(content: bytes) -> str:
    """Parse PDF file and extract text"""
    try:
        import pypdf
    except ImportError as e:
        raise ValueError("PDF parsing requires pypdf package") from e

    try:
        pdf_reader = pypdf.PdfReader(io.BytesIO(content))
        text_parts = [page.extract_text() for page in pdf_reader.pages if page.extract_text()]
        return "\n\n".join(text_parts)

    except Exception as e:
        logger.error("Failed to parse PDF: %s", e)
        raise ValueError(f"Failed to parse PDF file: {e}") from e


async def _parse_docx_file(content: bytes, file_ext: str) -> str:
    """Parse Word document and extract text"""
    if file_ext == ".doc":
        raise ValueError("Legacy .doc format is not supported. Please convert to .docx format.")

    try:
        import docx
    except ImportError as e:
        raise ValueError("DOCX parsing requires python-docx package") from e

    try:
        doc = docx.Document(io.BytesIO(content))
        text_parts = []

        # Extract paragraphs
        text_parts.extend(p.text for p in doc.paragraphs if p.text.strip())

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        return "\n\n".join(text_parts)

    except Exception as e:
        logger.error("Failed to parse DOCX: %s", e)
        raise ValueError(f"Failed to parse DOCX file: {e}") from e
